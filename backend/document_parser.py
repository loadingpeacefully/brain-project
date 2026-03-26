"""Document parsing — converts various file types to clean text."""
import io
import re


def parse_document(filename: str, content: bytes) -> str:
    """Parse a document from bytes into clean text."""
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return parse_pdf(content)
    elif name_lower.endswith(".docx"):
        return parse_docx(content)
    elif name_lower.endswith((".md", ".markdown")):
        return parse_markdown(content.decode("utf-8", errors="replace"))
    elif name_lower.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    elif name_lower.endswith(".csv"):
        return parse_csv(content)
    else:
        # Try as UTF-8 text
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {filename}")


def parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")
        return "\n\n".join(pages)
    except ImportError:
        raise ValueError("pypdf not installed. Run: pip install pypdf")


def parse_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name
            if style.startswith("Heading"):
                level = style[-1] if style[-1].isdigit() else "1"
                parts.append("#" * int(level) + " " + para.text)
            else:
                parts.append(para.text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")


def parse_markdown(text: str) -> str:
    # Remove HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Convert links to text only
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove image syntax
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    return text.strip()


def parse_csv(content: bytes) -> str:
    import csv
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    lines = [" | ".join(row) for row in rows]
    return "\n".join(lines)
